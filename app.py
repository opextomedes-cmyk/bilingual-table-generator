"""
=============================================================
BILINGUAL TABLE GENERATOR — Tomedes OPEX DTP Team
Powered by Claude AI + Streamlit
Supports: PDF files + Multiple Images (JPG, PNG, WEBP, GIF)
=============================================================
"""

import streamlit as st
import fitz  # PyMuPDF
import anthropic
import json
import re
import io
import os
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Page config ───────────────────────────────────────────────
st.set_page_config(
    page_title="Bilingual Table Generator",
    page_icon="📄",
    layout="centered"
)

# ── Styling ───────────────────────────────────────────────────
st.markdown("""
<style>
.main { background-color: #f0f4f8; }
.stApp { max-width: 900px; margin: 0 auto; }
h1 { color: #1E2761 !important; }
h2, h3 { color: #1E2761 !important; }
.stButton > button {
    background-color: #1E2761;
    color: white;
    font-weight: bold;
    border-radius: 8px;
    padding: 0.5rem 2rem;
    border: none;
    width: 100%;
}
.stButton > button:hover {
    background-color: #2d3a8a;
    color: white;
}
.tip-box {
    background: #fff8e1;
    border-left: 4px solid #ffc107;
    padding: 12px 16px;
    border-radius: 6px;
    font-size: 13px;
    color: #5d4037;
    margin: 10px 0;
}
.success-box {
    background: #f0fff4;
    border: 2px solid #a5d6a7;
    padding: 16px;
    border-radius: 10px;
    text-align: center;
    margin-top: 16px;
}
.header-box {
    background: #1E2761;
    color: white;
    padding: 20px 24px;
    border-radius: 12px;
    margin-bottom: 24px;
}
.mode-box {
    background: #e8eaf6;
    border-left: 4px solid #1E2761;
    padding: 10px 14px;
    border-radius: 6px;
    font-size: 13px;
    color: #1E2761;
    margin: 8px 0;
}
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────
st.markdown("""
<div class="header-box">
  <h2 style="color:white;margin:0;">📄 Bilingual Table Generator</h2>
  <p style="color:#CADCFC;margin:4px 0 0;">Tomedes OPEX DTP Team — Upload PDF or Images + Template → Get Bilingual .docx</p>
</div>
""", unsafe_allow_html=True)

# ── XML helpers for DOCX ──────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="B0BEC5", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

def set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# ── PDF → base64 for Claude ───────────────────────────────────
def pdf_to_base64(pdf_bytes):
    return base64.b64encode(pdf_bytes).decode('utf-8')

def image_to_base64(image_bytes):
    return base64.b64encode(image_bytes).decode('utf-8')

def get_image_media_type(filename):
    ext = filename.lower().split('.')[-1]
    mapping = {
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'png': 'image/png',
        'gif': 'image/gif',
        'webp': 'image/webp',
    }
    return mapping.get(ext, 'image/jpeg')

# ── Shared extraction prompt ──────────────────────────────────
EXTRACTION_PROMPT = """Extract ALL text from this document/image. Return ONLY raw JSON — no markdown, no explanation, no code fences.

CRITICAL RULES:
- DO NOT transcribe, retype, paraphrase, or summarize — extract EXACTLY as written
- Preserve bold, italic, underline indicators where detectable
- Keep sentences COMPLETE — never break mid-sentence
- Top-to-bottom, left-to-right reading order per page
- Non-text elements (stamps, images, seals, signatures, diagrams): use type "non_extractable"
- Each paragraph / heading / bullet / numbered item = one separate element
- Include ALL text: headers, footers, labels, captions, field names, field values
- If a word appears bold, set "bold": true. If italic, set "italic": true. If underlined, set "underline": true.

Return this exact JSON format:
{"pages":[{"pageNumber":1,"elements":[{"type":"heading|paragraph|bullet|numbered|caption|label|non_extractable","text":"exact extracted text","bold":false,"italic":false,"underline":false}]}],"totalPages":1}"""

# ── Claude extraction: PDF ────────────────────────────────────
def extract_pdf_with_claude(pdf_bytes, api_key):
    client = anthropic.Anthropic(api_key=api_key)
    pdf_b64 = pdf_to_base64(pdf_bytes)

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_b64
                    }
                },
                {
                    "type": "text",
                    "text": EXTRACTION_PROMPT
                }
            ]
        }]
    )

    return parse_claude_response(response)

# ── Claude extraction: Single Image ──────────────────────────
def extract_image_with_claude(image_bytes, filename, page_number, api_key):
    client = anthropic.Anthropic(api_key=api_key)
    img_b64 = image_to_base64(image_bytes)
    media_type = get_image_media_type(filename)

    prompt = EXTRACTION_PROMPT + f"\n\nThis is image {page_number} (filename: {filename}). Set pageNumber to {page_number} in your response."

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": img_b64
                    }
                },
                {
                    "type": "text",
                    "text": prompt
                }
            ]
        }]
    )

    result = parse_claude_response(response)

    # Ensure page number matches index
    if result.get("pages"):
        result["pages"][0]["pageNumber"] = page_number
        result["pages"][0]["sourceFile"] = filename
    else:
        result["pages"] = [{
            "pageNumber": page_number,
            "sourceFile": filename,
            "elements": [{"type": "non_extractable", "text": f"[NON-EXTRACTABLE CONTENT — Image {page_number}: {filename}]", "bold": False, "italic": False, "underline": False}]
        }]

    return result

# ── Parse Claude response ─────────────────────────────────────
def parse_claude_response(response):
    raw = response.content[0].text if response.content else ""
    raw = raw.replace("```json", "").replace("```", "").strip()
    try:
        return json.loads(raw)
    except Exception:
        match = re.search(r'\{[\s\S]*\}', raw)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
    return {"pages": [], "totalPages": 0}

# ── Build DOCX ────────────────────────────────────────────────
def build_docx(pages, source_lang, target_lang, font_size, input_mode="pdf"):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(font_size)

    COL_W = 4050
    BLUE = "2E74B5"
    ICE = "D9E2F3"
    NAVY = "1E2761"
    WHITE = "FFFFFF"
    SHADE = "F4F6FB"

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    def add_col_header():
        row = table.add_row()
        for i, lang_label in enumerate([source_lang, target_lang]):
            c = row.cells[i]
            set_cell_bg(c, BLUE)
            set_cell_borders(c, color=BLUE, size="6")
            set_cell_margins(c, top=100, bottom=100, left=150, right=150)
            set_col_width(c, COL_W)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lang_label)
            r.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def add_page_header(label):
        row = table.add_row()
        for i in range(2):
            c = row.cells[i]
            set_cell_bg(c, ICE)
            set_cell_borders(c, color="1E2761", size="6")
            set_cell_margins(c, top=100, bottom=100, left=150, right=150)
            set_col_width(c, COL_W)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)

    def add_content_row(el, shade, page_num):
        row = table.add_row()
        src = row.cells[0]
        tgt = row.cells[1]
        bg = SHADE if shade else WHITE
        for c in [src, tgt]:
            set_cell_bg(c, bg)
            set_cell_borders(c, color="B0BEC5", size="2")
            set_cell_margins(c)
            set_col_width(c, COL_W)

        p = src.paragraphs[0]
        txt = el.get("text", "")

        if el.get("type") == "non_extractable":
            txt = txt if txt.startswith("[NON-EXTRACTABLE") else f"[NON-EXTRACTABLE CONTENT — Page {page_num}]"
            r = p.add_run(txt)
            r.italic = True
            r.font.name = 'Calibri'
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        else:
            r = p.add_run(txt)
            r.font.name = 'Calibri'
            r.font.size = Pt(font_size)
            r.bold = el.get("bold", False)
            r.italic = el.get("italic", False)
            if el.get("underline", False):
                r.underline = True

        tp = tgt.paragraphs[0]
        tr = tp.add_run("")
        tr.font.name = 'Calibri'
        tr.font.size = Pt(font_size)

    add_col_header()
    total = 0

    for page in pages:
        page_num = page.get("pageNumber", 1)
        elements = page.get("elements", [])
        source_file = page.get("sourceFile", "")

        if input_mode == "images" and source_file:
            header_label = f"Image {page_num} — {source_file}"
        else:
            header_label = f"Page {page_num}"

        add_page_header(header_label)

        if not elements:
            fake_el = {"type": "paragraph", "text": "[No extractable content on this page]", "italic": True}
            add_content_row(fake_el, False, page_num)
        else:
            for idx, el in enumerate(elements):
                add_content_row(el, idx % 2 == 0, page_num)
                total += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, total


# ═════════════════════════════════════════════════════════════
# UI
# ═════════════════════════════════════════════════════════════

# ── Step 1: API Key ───────────────────────────────────────────
st.markdown("### 🔑 Step 1 — Enter Your Claude API Key")
st.markdown("""
<div class="tip-box">
Your API key is used only for this session and is never stored.
Get yours free at <a href="https://console.anthropic.com" target="_blank">console.anthropic.com</a>
</div>
""", unsafe_allow_html=True)

api_key = st.text_input(
    "Claude API Key",
    type="password",
    placeholder="sk-ant-...",
    label_visibility="collapsed"
)

st.divider()

# ── Step 2: Input Mode ────────────────────────────────────────
st.markdown("### 📁 Step 2 — Choose Input Type & Upload Files")

input_mode = st.radio(
    "What are you uploading as source?",
    options=["PDF", "Images (JPG, PNG, WEBP, GIF)"],
    horizontal=True
)

col1, col2 = st.columns(2)

pdf_file = None
image_files = []

with col1:
    if input_mode == "PDF":
        st.markdown("**Source PDF**")
        pdf_file = st.file_uploader(
            "Upload PDF",
            type=["pdf"],
            label_visibility="collapsed",
            key="pdf_upload"
        )
        if pdf_file:
            st.success(f"✓ {pdf_file.name}")
    else:
        st.markdown("**Source Images** *(upload one or more)*")
        image_files = st.file_uploader(
            "Upload Images",
            type=["jpg", "jpeg", "png", "gif", "webp"],
            accept_multiple_files=True,
            label_visibility="collapsed",
            key="img_upload"
        )
        if image_files:
            st.success(f"✓ {len(image_files)} image(s) uploaded")
            for f in image_files:
                st.caption(f"  • {f.name}")

with col2:
    st.markdown("**Word Template (.docx)**")
    tpl_file = st.file_uploader(
        "Upload Template",
        type=["docx"],
        label_visibility="collapsed",
        key="tpl_upload"
    )
    if tpl_file:
        st.success(f"✓ {tpl_file.name}")

if input_mode == "Images (JPG, PNG, WEBP, GIF)" and image_files:
    st.markdown(f"""
    <div class="mode-box">
    📸 <strong>Image mode:</strong> {len(image_files)} image(s) will be processed in upload order.
    Each image = one section in the bilingual table (Image 1, Image 2, …).
    Tip: rename your files with a number prefix (e.g. 01_page.jpg) before uploading to control order.
    </div>
    """, unsafe_allow_html=True)

st.divider()

# ── Step 3: Settings ──────────────────────────────────────────
st.markdown("### ⚙️ Step 3 — Settings")

col3, col4, col5 = st.columns(3)
with col3:
    source_lang = st.text_input("Source Language", value="English", placeholder="e.g. English")
with col4:
    target_lang = st.text_input("Target Language", value="", placeholder="e.g. French")
with col5:
    font_size = st.selectbox("Font Size", options=[10, 11, 12], index=0, format_func=lambda x: f"{x}pt")

st.markdown("""
<div class="tip-box">
⚠️ <strong>No transcription:</strong> All text is extracted directly from the source by Claude AI.
Non-extractable content (stamps, seals, diagrams, handwriting) will be flagged as
<code>[NON-EXTRACTABLE CONTENT — Page X]</code>.
</div>
""", unsafe_allow_html=True)

st.divider()

# ── Step 4: Generate ──────────────────────────────────────────
st.markdown("### 🚀 Step 4 — Generate")

has_source = pdf_file if input_mode == "PDF" else (len(image_files) > 0 if image_files else False)
ready = api_key and has_source and tpl_file

if not ready:
    missing = []
    if not api_key:
        missing.append("API Key")
    if not has_source:
        missing.append("PDF file" if input_mode == "PDF" else "at least one image")
    if not tpl_file:
        missing.append("Word template")
    st.info(f"Still needed: {', '.join(missing)}")

generate_btn = st.button("⚡ Generate Bilingual Table", disabled=not ready)

if generate_btn and ready:
    try:
        all_pages = []

        if input_mode == "PDF":
            with st.spinner("📤 Sending PDF to Claude AI for extraction..."):
                pdf_bytes = pdf_file.read()
                extracted = extract_pdf_with_claude(pdf_bytes, api_key)
                all_pages = extracted.get("pages", [])
                total_pages = len(all_pages)

            if total_pages == 0:
                st.error("Could not extract content from the PDF. Please check the file and try again.")
                st.stop()

            st.success(f"✓ Extracted {total_pages} page(s) from PDF")
            with st.expander("📋 Extraction summary"):
                for p in all_pages:
                    count = len(p.get("elements", []))
                    st.write(f"• Page {p.get('pageNumber', '?')}: {count} element(s)")

        else:
            # Images mode — process each image individually
            progress = st.progress(0, text="Processing images...")
            total_images = len(image_files)

            for idx, img_file in enumerate(image_files):
                page_number = idx + 1
                progress.progress(idx / total_images, text=f"Processing image {page_number}/{total_images}: {img_file.name}")

                with st.spinner(f"📤 Extracting text from image {page_number}: {img_file.name}..."):
                    img_bytes = img_file.read()
                    result = extract_image_with_claude(img_bytes, img_file.name, page_number, api_key)
                    if result.get("pages"):
                        all_pages.extend(result["pages"])

            progress.progress(1.0, text="All images processed!")

            if not all_pages:
                st.error("Could not extract content from the images. Please check the files and try again.")
                st.stop()

            st.success(f"✓ Extracted content from {len(all_pages)} image(s)")
            with st.expander("📋 Extraction summary"):
                for p in all_pages:
                    count = len(p.get("elements", []))
                    label = p.get("sourceFile", f"Image {p.get('pageNumber', '?')}")
                    st.write(f"• Image {p.get('pageNumber', '?')} ({label}): {count} element(s)")

        # Build DOCX
        with st.spinner("📝 Building bilingual Word document..."):
            mode_key = "images" if input_mode != "PDF" else "pdf"
            doc_buf, total_rows = build_docx(
                all_pages,
                source_lang,
                target_lang or "Translation",
                font_size,
                input_mode=mode_key
            )

        # Output name
        if input_mode == "PDF":
            base_name = pdf_file.name.replace(".pdf", "").replace(".PDF", "")
        else:
            base_name = "images_combined"

        output_name = f"{base_name}_bilingual.docx"

        st.markdown(f"""
        <div class="success-box">
          <h3 style="color:#2e7d32;">✅ Bilingual Table Ready!</h3>
          <p style="color:#546e7a;">
            {len(all_pages)} section(s) &nbsp;·&nbsp;
            {total_rows} content rows &nbsp;·&nbsp;
            Calibri {font_size}pt &nbsp;·&nbsp;
            Target column ready for translation
          </p>
        </div>
        """, unsafe_allow_html=True)

        st.download_button(
            label="⬇️ Download Bilingual .docx",
            data=doc_buf,
            file_name=output_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except anthropic.AuthenticationError:
        st.error("❌ Invalid API key. Please check your Claude API key and try again.")
    except anthropic.RateLimitError:
        st.error("❌ Rate limit reached. Please wait a moment and try again.")
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
        st.exception(e)

# ── Footer ────────────────────────────────────────────────────
st.divider()
st.markdown("""
<p style="text-align:center;font-size:12px;color:#90a4ae;">
Tomedes OPEX DTP Team &nbsp;·&nbsp; Bilingual Table Generator &nbsp;·&nbsp; Powered by Claude AI
</p>
""", unsafe_allow_html=True)
